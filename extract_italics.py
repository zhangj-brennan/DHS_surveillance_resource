from docx import Document
from docx.text.run import Run
from docx.oxml.ns import qn
import json

DOCX_FILE = "DHS interior surveillance resource - FINAL pre legal review (2).docx"
OUTPUT_JSON = "italic_phrases.json"

doc = Document(DOCX_FILE)


def is_run_italic(run):
    # Direct formatting
    if run.italic is True:
        return True

    # Character-style formatting
    if run.style is not None and run.style.font.italic is True:
        return True

    # Underlying XML formatting
    rpr = run._element.rPr

    if rpr is not None:
        italic_element = rpr.find(qn("w:i"))

        if italic_element is not None:
            value = italic_element.get(qn("w:val"))

            # <w:i/> means italic.
            # <w:i w:val="0"/> means not italic.
            return value not in ("0", "false", "False", "off")

    return False


def get_paragraph_runs_in_order(paragraph):
    """
    Returns all runs in their original order, including runs inside hyperlinks.
    """

    for child in paragraph._p.iterchildren():
        # Normal run
        if child.tag == qn("w:r"):
            yield Run(child, paragraph)

        # Hyperlink containing one or more runs
        elif child.tag == qn("w:hyperlink"):
            for run_element in child.findall(qn("w:r")):
                yield Run(run_element, paragraph)


def extract_italic_phrases(paragraph):
    phrases = []
    current_phrase = ""

    for run in get_paragraph_runs_in_order(paragraph):
        if is_run_italic(run):
            current_phrase += run.text
        else:
            phrase = current_phrase.strip()

            if phrase:
                phrases.append(phrase)

            current_phrase = ""

    final_phrase = current_phrase.strip()

    if final_phrase:
        phrases.append(final_phrase)

    return phrases


italic_phrases = {}

# Standard document paragraphs
for paragraph in doc.paragraphs:
    for phrase in extract_italic_phrases(paragraph):
        italic_phrases.setdefault(phrase, {})

# Paragraphs inside tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for phrase in extract_italic_phrases(paragraph):
                    italic_phrases.setdefault(phrase, {})

# Headers
for section in doc.sections:
    for paragraph in section.header.paragraphs:
        for phrase in extract_italic_phrases(paragraph):
            italic_phrases.setdefault(phrase, {})

# Footers
for section in doc.sections:
    for paragraph in section.footer.paragraphs:
        for phrase in extract_italic_phrases(paragraph):
            italic_phrases.setdefault(phrase, {})

with open(OUTPUT_JSON, "w", encoding="utf-8") as output_file:
    json.dump(
        italic_phrases,
        output_file,
        indent=2,
        ensure_ascii=False
    )

print(f"Found {len(italic_phrases)} unique italic phrases.")
print(f"Saved to: {OUTPUT_JSON}")